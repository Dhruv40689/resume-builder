
import re
import io
import warnings
import logging


logging.getLogger("pdfminer").setLevel(logging.ERROR)
warnings.filterwarnings("ignore", message=".*FontBBox.*")
warnings.filterwarnings("ignore", message=".*CryptographyDeprecation.*")
warnings.filterwarnings("ignore", message=".*float.*")


class ResumeParser:

    def __init__(self):
        self.email_re    = re.compile(r'\b[\w._%+-]+@[\w.-]+\.[a-zA-Z]{2,}\b')
        self.phone_re    = re.compile(r'[\+\(]?[\d][\d\s.\-\(\)]{7,}[\d]')
        self.linkedin_re = re.compile(r'linkedin\.com/in/[\w\-]+', re.I)
        self.github_re   = re.compile(r'github\.com/[\w\-/]+', re.I)

        
        self.style_section_map = {
            'heading 1': None,   
            'heading 2': None,
            'heading 3': None,
            'title':     'name',
            'subtitle':  'contact',
        }

        self.section_keywords = {
            'summary':        ['summary', 'objective', 'profile', 'about', 'overview', 'professional summary'],
            'experience':     ['experience', 'employment', 'work history', 'work experience', 'career'],
            'education':      ['education', 'academic', 'qualification', 'degree'],
            'skills':         ['skills', 'technical skills', 'competencies', 'expertise', 'technologies', 'skills & expertise'],
            'projects':       ['projects', 'portfolio', 'achievements', 'projects & achievements'],
            'certifications': ['certifications', 'certificates', 'licenses', 'credentials'],
        }

 
    def parse_pdf(self, file_bytes: bytes):
        text = self._extract_pdf_text(file_bytes)
        return self._extract_from_text(text), text

    def parse_docx(self, file_bytes: bytes):
       
        try:
            from docx import Document
            doc = Document(io.BytesIO(file_bytes))
            data, full_text = self._extract_from_docx(doc)
            return data, full_text
        except Exception as e:
            return {'name': '', 'email': '', 'summary': '', 'skills': []}, str(e)

    

    def _extract_from_docx(self, doc):
       
        data = {
            'name': '', 'email': '', 'phone': '', 'linkedin': '',
            'github': '', 'website': '', 'location': '', 'summary': '',
            'skills': [], 'education_text': '', 'experience_text': '',
            'projects_text': '', 'certifications': [],
        }

        sections = {}         
        current_section = 'header'
        current_lines   = []
        exp_entries     = []   
        current_exp     = None

        full_text_parts = []

        for para in doc.paragraphs:
            text  = para.text.strip()
            style = para.style.name.lower() if para.style else 'normal'

            if not text:
                continue

            full_text_parts.append(text)

           
            if 'title' in style and 'subtitle' not in style:
                data['name'] = text.title() if text.isupper() else text
                continue

           
            if 'subtitle' in style:
                self._parse_contact_line(text, data)
                continue

            
            if 'heading 1' in style:
                sections[current_section] = current_lines
                current_section = self._classify_section(text)
                current_lines   = []
                
                if current_exp and current_section != 'experience':
                    exp_entries.append(current_exp)
                    current_exp = None
                continue

           
            if 'heading 2' in style and current_section == 'experience':
                
                if current_exp:
                    exp_entries.append(current_exp)
                current_exp = {'duration': text, 'title': '', 'company': '', 'responsibilities': []}
                continue

            
            if 'heading 3' in style and current_section == 'experience':
                if current_exp is None:
                    current_exp = {'duration': '', 'title': '', 'company': '', 'responsibilities': []}
                
                if '|' in text:
                    parts = text.split('|', 1)
                    current_exp['title']   = parts[0].strip()
                    current_exp['company'] = parts[1].strip()
                else:
                    current_exp['title'] = text.strip()
                continue

          
            if current_section == 'experience' and current_exp is not None:
                current_exp['responsibilities'].append(text)
            else:
                current_lines.append(text)

       
        sections[current_section] = current_lines
        if current_exp:
            exp_entries.append(current_exp)

        
        exp_parts = []
        for exp in exp_entries:
            header = f"{exp['title']} | {exp['company']} ({exp['duration']})" if exp['company'] else f"{exp['title']} ({exp['duration']})"
            lines  = [header] + ['• ' + r for r in exp['responsibilities']]
            exp_parts.append('\n'.join(lines))
        data['experience_text']  = '\n\n'.join(exp_parts)
        data['experience_entries'] = exp_entries

        
        for sec_type, lines in sections.items():
            body = '\n'.join(lines).strip()
            if not body:
                continue
            if sec_type == 'summary':
                data['summary'] = body
            elif sec_type == 'education':
                data['education_text'] = body
            elif sec_type == 'projects':
                data['projects_text'] = body
            elif sec_type == 'skills':
                data['skills'] = self._parse_skills(body)
            elif sec_type == 'certifications':
                data['certifications'] = [l for l in body.split('\n') if l.strip()]

     
        if not data['skills']:
            data['skills'] = self._fallback_skills('\n'.join(full_text_parts))

       
        if not data['name']:
            lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            data['name'] = self._extract_name_from_lines(lines, data)

        full_text = '\n'.join(full_text_parts)
        return data, full_text

    def _classify_section(self, heading_text: str) -> str:
        ht = heading_text.lower().strip()
        for sec_type, keywords in self.section_keywords.items():
            if any(kw in ht for kw in keywords):
                return sec_type
        return heading_text.lower()

    def _parse_contact_line(self, line: str, data: dict):
        
        em = self.email_re.search(line)
        if em:
            data['email'] = em.group()

        ph = self.phone_re.search(line)
        if ph:
            digits = re.sub(r'[^\d]', '', ph.group())
            if len(digits) > 10:
                digits = digits[-10:]
            if len(digits) >= 10:
                data['phone'] = digits

        li = self.linkedin_re.search(line)
        if li:
            data['linkedin'] = li.group()

        gh = self.github_re.search(line)
        if gh:
            data['github']  = gh.group()
            data['website'] = gh.group()

        
        cleaned = line
        for pattern in [self.email_re, self.phone_re, self.linkedin_re, self.github_re]:
            cleaned = pattern.sub('', cleaned)
        cleaned = re.sub(r'[•|,]', ' ', cleaned)
       
        cleaned = re.sub(r'\+?\d[\d\s\-\(\)]{6,}', '', cleaned)
        cleaned = ' '.join(cleaned.split())
        if cleaned and len(cleaned) > 2:
            data['location'] = cleaned.strip()

    

    def _extract_pdf_text(self, file_bytes: bytes) -> str:
        try:
            import pdfplumber
           
            with warnings.catch_warnings():
                warnings.simplefilter("ignore")
                with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                    return "\n".join(p.extract_text() or "" for p in pdf.pages)
        except ImportError:
            pass
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            return "\n".join(p.extract_text() or "" for p in reader.pages)
        except Exception as e:
            return f"(PDF parse error: {e})"

   

    def _extract_from_text(self, text: str) -> dict:
        data = {
            'name': '', 'email': '', 'phone': '', 'linkedin': '',
            'github': '', 'website': '', 'location': '', 'summary': '',
            'skills': [], 'education_text': '', 'experience_text': '',
            'projects_text': '', 'certifications': [],
        }
        lines = [l.strip() for l in text.split('\n') if l.strip()]

        em = self.email_re.search(text)
        if em: data['email'] = em.group()

        ph = self.phone_re.search(text)
        if ph:
            digits = re.sub(r'[^\d]', '', ph.group())
            if len(digits) > 10: digits = digits[-10:]
            if len(digits) >= 10: data['phone'] = digits

        li = self.linkedin_re.search(text)
        if li: data['linkedin'] = li.group()

        gh = self.github_re.search(text)
        if gh:
            data['github']  = gh.group()
            data['website'] = gh.group()

        
        for pat in [
            r'\b([A-Z][a-z]+(?:\s[A-Z][a-z]+)?,\s*(?:India|Maharashtra|[A-Z]{2}))\b',
            r'\b(Bhayadar|Mumbai|Delhi|Vasai|Thane|Pune|Bangalore|Bengaluru|Hyderabad)\b',
        ]:
            m = re.search(pat, text)
            if m: data['location'] = m.group(); break

        data['name'] = self._extract_name_from_lines(lines, data)
        sections     = self._split_sections_text(lines)

        for sec_type, keywords in self.section_keywords.items():
            for sec_name, content in sections.items():
                if any(kw in sec_name.lower() for kw in keywords):
                    body = '\n'.join(content).strip()
                    if   sec_type == 'summary':        data['summary']        = body
                    elif sec_type == 'experience':     data['experience_text']= body
                    elif sec_type == 'education':      data['education_text'] = body
                    elif sec_type == 'projects':       data['projects_text']  = body
                    elif sec_type == 'skills':         data['skills']         = self._parse_skills(body)
                    elif sec_type == 'certifications': data['certifications'] = [l for l in body.split('\n') if l.strip()]

        if not data['skills']:
            data['skills'] = self._fallback_skills(text)

        return data

    def _split_sections_text(self, lines: list) -> dict:
        all_kws = [kw for kws in self.section_keywords.values() for kw in kws]

        def is_header(line):
            ll = line.lower().strip().rstrip(':')
            if ll in all_kws: return True
            if line.isupper() and 1 < len(line.split()) <= 5: return True
            if line.endswith(':') and len(line) < 40: return True
            return False

        sections = {}
        current  = 'header'
        content  = []
        for line in lines:
            if is_header(line):
                sections[current] = content
                current = line.lower().strip().rstrip(':')
                content = []
            else:
                content.append(line)
        sections[current] = content
        return sections

    

    def _extract_name_from_lines(self, lines: list, data: dict) -> str:
        skip_words  = {'resume','cv','profile','summary','education','experience',
                       'skills','projects','certifications','contact','objective','professional'}
        contact_sig = {'@','github','linkedin','http','www','.com','.in','+91'}

        def is_contact(line):
            ll = line.lower()
            return any(s in ll for s in contact_sig) or '|' in line or '•' in line

        def is_name(token: str) -> bool:
            words = token.strip().split()
            if not (2 <= len(words) <= 4): return False
            if any(w.lower() in skip_words for w in words): return False
            return all(re.match(r"^[A-Za-z\-']+$", w) for w in words)
 
        
        for line in lines[:12]:
            if not is_contact(line) and is_name(line):
                return line.title() if line.isupper() else line

        for line in lines[:5]:
            cleaned = line
            cleaned = self.email_re.sub('', cleaned)
            cleaned = re.sub(r'https?://\S+', '', cleaned)
            cleaned = self.github_re.sub('', cleaned)
            cleaned = self.linkedin_re.sub('', cleaned)
            cleaned = re.sub(r'\b(Engineering|Location|Mumbai|Delhi|Bangalore|Vasai|Thane|Bhayadar|India)\b', '', cleaned, flags=re.I)
            cleaned = re.sub(r'[|•,\d\+\(\)\-/]', ' ', cleaned)
            cleaned = ' '.join(cleaned.split())
            if is_name(cleaned):
                return cleaned.title()

       
        first = lines[0] if lines else ''
        last_seg = first.split('|')[-1].strip()
        last_seg = re.sub(r'\S+\.\S+', ' ', last_seg)
        words = last_seg.split()
        for n in (2, 3):
            candidate = ' '.join(words[-n:])
            if is_name(candidate):
                return candidate.title()

        return first[:50]



    def _parse_skills(self, body: str) -> list:
        normalized = re.sub(r'[•·|▸▪\n]', ',', body)
        normalized = re.sub(r'[A-Za-z &/()]+:\s*', '', normalized)
        parts  = [p.strip() for p in normalized.split(',')]
        skills = []
        bad_starts = ('demonstrating', 'i am', 'with', 'and ', 'the ', 'other basics')
        for p in parts:
            p = p.strip(' -–:•·')
            if (1 < len(p) < 50 and
                    not any(p.lower().startswith(b) for b in bad_starts) and
                    len(p.split()) <= 5):
                skills.append(p)
        return list(dict.fromkeys(skills))

    def _fallback_skills(self, text: str) -> list:
        known = [
            'Python','JavaScript','TypeScript','Java','C++','C#','Go',
            'React','Angular','Vue','Next.js','Node.js','Express',
            'Flutter','Dart','Android','iOS','Swift','Kotlin',
            'Django','Flask','FastAPI','Spring Boot',
            'SQL','PostgreSQL','MySQL','MongoDB','Redis','Firebase',
            'AWS','GCP','Azure','Docker','Kubernetes','Linux','Git',
            'TensorFlow','PyTorch','Scikit-learn','Pandas','NumPy',
            'Machine Learning','Deep Learning','NLP','LLM','Generative AI',
            'LangChain','Hugging Face','RAG',
            'REST API','GraphQL','Microservices','CI/CD','Agile',
            'HTML','CSS','Tailwind','Bootstrap',
        ]
        tl = text.lower()
        return [k for k in known if k.lower() in tl]