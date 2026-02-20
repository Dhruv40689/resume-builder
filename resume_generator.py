import io
from typing import Dict, Optional


class ResumeGenerator:

    TEMPLATES = {
        "Classic Professional": {
            "header_color": "2C3E50",
            "accent_color": "2980B9",
            "secondary_color": "ECF0F1",
            "font_name": "Arial",
            "header_style": "traditional"
        },
        "Modern Minimalist": {
            "header_color": "1A1A2E",
            "accent_color": "16213E",
            "secondary_color": "E8F4FD",
            "font_name": "Calibri",
            "header_style": "modern"
        },
        "Executive Bold": {
            "header_color": "1B2631",
            "accent_color": "C0392B",
            "secondary_color": "FDFEFE",
            "font_name": "Georgia",
            "header_style": "executive"
        }
    }

    def generate_docx(self, resume_data: Dict, template_name: str = "Classic Professional") -> bytes:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import re

        template = self.TEMPLATES.get(template_name, self.TEMPLATES["Classic Professional"])
        doc = Document()

        for section in doc.sections:
            section.top_margin = Inches(0.7)
            section.bottom_margin = Inches(0.7)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)

        def hex_to_rgb(hex_color: str):
            h = hex_color.lstrip('#')
            return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))

        accent_rgb = hex_to_rgb(template['accent_color'])
        header_rgb = hex_to_rgb(template['header_color'])

        def add_horizontal_line(doc, color_hex="2980B9", thickness=15):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), str(thickness))
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), color_hex)
            pBdr.append(bottom)
            pPr.append(pBdr)
            return p

        def add_section_header(doc, title: str, template: dict):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(title.upper())
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = template['font_name']
            run.font.color.rgb = RGBColor(*hex_to_rgb(template['accent_color']))
            add_horizontal_line(doc, template['accent_color'], 12)
            return p

        def add_paragraph(doc, text: str, font_size: int = 10, bold: bool = False,
                          italic: bool = False, indent: bool = False, space_after: int = 2):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(space_after)
            if indent:
                p.paragraph_format.left_indent = Inches(0.15)
            run = p.add_run(text)
            run.font.size = Pt(font_size)
            run.font.name = template['font_name']
            run.bold = bold
            run.italic = italic
            return p

        name = resume_data.get('name', 'Your Name')
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_para.paragraph_format.space_before = Pt(0)
        name_para.paragraph_format.space_after = Pt(7)
        name_run = name_para.add_run(name)
        name_run.bold = True
        name_run.font.size = Pt(24)
        name_run.font.name = template['font_name']
        name_run.font.color.rgb = RGBColor(*header_rgb)

        contact_parts = []
        for field in ['email', 'phone', 'location', 'linkedin', 'website']:
            if resume_data.get(field):
                contact_parts.append(resume_data[field])

        if contact_parts:
            contact_para = doc.add_paragraph()
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_para.paragraph_format.space_before = Pt(4)
            contact_para.paragraph_format.space_after = Pt(8)
            contact_run = contact_para.add_run(" | ".join(contact_parts))
            contact_run.font.size = Pt(10.5)
            contact_run.font.name = template['font_name']

        add_horizontal_line(doc, template['accent_color'], 20)

        summary = resume_data.get('summary', '')
        if summary:
            add_section_header(doc, "Professional Summary", template)
            p = doc.add_paragraph()
            run = p.add_run(summary)
            run.font.size = Pt(10)
            run.font.name = template['font_name']

        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        return doc_buffer.read()

    def generate_pdf(self, resume_data: Dict, template_name: str = "Classic Professional") -> bytes:
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.lib import colors
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
            from reportlab.lib.enums import TA_CENTER

            template = self.TEMPLATES.get(template_name, self.TEMPLATES["Classic Professional"])

            def hex_to_color(hex_str):
                h = hex_str.lstrip('#')
                r, g, b = int(h[0:2], 16)/255, int(h[2:4], 16)/255, int(h[4:6], 16)/255
                return colors.Color(r, g, b)

            accent_color = hex_to_color(template['accent_color'])
            header_color = hex_to_color(template['header_color'])

            buffer = io.BytesIO()
            doc = SimpleDocTemplate(
                buffer,
                pagesize=letter,
                leftMargin=0.75*inch,
                rightMargin=0.75*inch,
                topMargin=0.7*inch,
                bottomMargin=0.7*inch
            )

            styles = getSampleStyleSheet()

            name_style = ParagraphStyle('NameStyle',
                fontName='Helvetica-Bold',
                fontSize=22,
                textColor=header_color,
                alignment=TA_CENTER,
                spaceAfter=4
            )

            story = []
            story.append(Paragraph(resume_data.get('name', 'Your Name'), name_style))
            story.append(HRFlowable(width="100%", thickness=2, color=accent_color, spaceAfter=8))

            doc.build(story)
            buffer.seek(0)
            return buffer.read()

        except ImportError:
            raise Exception("ReportLab not installed. Install with: pip install reportlab")
        except Exception as e:
            raise Exception(f"PDF generation failed: {str(e)}")